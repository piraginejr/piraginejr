#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Guia_Implementacao_Email_Moderno_Microsoft_Graph_Power_Church.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def base_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 18, 10),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 14, 7),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    return doc


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Guia de Implementacao do Envio Moderno de E-mail\nMicrosoft Graph + Power Church Django")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    r2 = p2.add_run(
        "Documento passo a passo para iniciantes. "
        "Este guia cobre desde a configuracao da Microsoft 365 ate o teste final do envio automatico de recibos."
    )
    r2.font.size = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta.add_run("Projeto: ").bold = True
    meta.add_run("Power Church Django\n")
    meta.add_run("Objetivo: ").bold = True
    meta.add_run("Enviar recibos com autenticacao moderna, sem depender de SMTP AUTH basico\n")
    meta.add_run("Conta de envio prevista: ").bold = True
    meta.add_run("recebimento@pibn.org.br")


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Power Church | Guia de Implementacao do E-mail Moderno"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Pagina ")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    add_page_number(footer)


def add_intro_table(doc: Document) -> None:
    doc.add_heading("1. Visao Geral", level=1)
    doc.add_paragraph(
        "A recomendacao deste projeto e usar Microsoft Graph com OAuth 2.0 em modo aplicativo "
        "(app-only). Isso evita dependencia de senha de mailbox e segue o caminho moderno da Microsoft."
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = ["Decisao", "Resumo"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_width(cell, 1.85 if idx == 0 else 4.65)
        set_cell_shading(cell, "E8EEF5")
    rows = [
        ("Metodo escolhido", "Microsoft Graph API com OAuth 2.0 e permissao Mail.Send"),
        ("Modo de autenticacao", "App registration no Microsoft Entra ID com client secret ou certificado"),
        ("Envio no sistema", "Django gera o PDF do recibo e envia um e-mail por vez pela API"),
        ("Remetente previsto", "recebimento@pibn.org.br"),
        ("Uso no projeto", "Recibos automáticos por evento e reenvio manual pelo operador"),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        set_cell_width(row[0], 1.85)
        set_cell_width(row[1], 4.65)


def add_prereqs(doc: Document) -> None:
    doc.add_heading("2. O que voce precisa antes de comecar", level=1)
    items = [
        "Acesso de administrador ao Microsoft 365 ou ajuda de quem tenha esse acesso.",
        "A conta de envio ja criada: recebimento@pibn.org.br.",
        "Acesso ao projeto Power Church Django e ao arquivo de configuracao de ambiente.",
        "Um terminal com PowerShell para a etapa do Exchange Online.",
        "Tempo para executar os testes com calma, sem pressa.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "Importante: para producao, o ideal e usar certificado em vez de client secret. "
        "Para a primeira implantacao, o client secret e mais facil e rapido."
    )


def add_checklist_table(doc: Document) -> None:
    doc.add_heading("3. Checklist rapido de informacoes", level=1)
    doc.add_paragraph("Antes de entrar no portal, tenha estes dados anotados:")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [2.2, 2.0, 2.3]
    for idx, text in enumerate(["Item", "Valor esperado", "Anotado?"]):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E8EEF5")
    data = [
        ("Tenant principal", "Seu tenant Microsoft 365", "[   ]"),
        ("Mailbox remetente", "recebimento@pibn.org.br", "[   ]"),
        ("Nome da app", "Power Church Mailer", "[   ]"),
        ("Ambiente", "Producao ou homologacao", "[   ]"),
        ("URL do projeto", "Servidor/local do Django", "[   ]"),
    ]
    for row_values in data:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value
            set_cell_width(row[idx], widths[idx])


def add_entra_steps(doc: Document) -> None:
    doc.add_heading("4. Criar o aplicativo no Microsoft Entra ID", level=1)
    doc.add_paragraph("Siga exatamente estes passos:")
    steps = [
        "Entre em https://entra.microsoft.com com uma conta administradora.",
        "No menu lateral, abra App registrations.",
        "Clique em New registration.",
        "No campo Name, escreva: Power Church Mailer.",
        "Escolha Accounts in this organizational directory only (Single tenant).",
        "Deixe Redirect URI vazio nesta fase, porque vamos usar app-only.",
        "Clique em Register.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph("Quando a tela abrir, anote estes dois campos:")
    for item in [
        "Application (client) ID",
        "Directory (tenant) ID",
    ]:
        doc.add_paragraph(item, style="List Bullet")


def add_secret_steps(doc: Document) -> None:
    doc.add_heading("5. Criar a credencial da aplicacao", level=1)
    doc.add_paragraph("Para um iniciante, o modo mais simples e usar client secret.")
    steps = [
        "Dentro da app recem-criada, entre em Certificates & secrets.",
        "Clique em New client secret.",
        "Em Description, escreva algo como: Power Church Django Producao.",
        "Escolha um prazo adequado. Para testes, pode ser curto; para producao, planeje renovacao.",
        "Clique em Add.",
        "Copie o campo Value imediatamente e guarde com seguranca. Depois ele nao aparece novamente.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "Guarde o valor em local seguro. Esse valor sera a variavel POWER_CHURCH_GRAPH_CLIENT_SECRET no Django."
    )


def add_permissions_steps(doc: Document) -> None:
    doc.add_heading("6. Dar permissao Mail.Send ao aplicativo", level=1)
    steps = [
        "Na app, abra API permissions.",
        "Clique em Add a permission.",
        "Escolha Microsoft Graph.",
        "Escolha Application permissions.",
        "Pesquise por Mail.Send.",
        "Marque Mail.Send.",
        "Clique em Add permissions.",
        "Depois clique em Grant admin consent for <tenant>.",
        "Confirme que a coluna Status mudou para Granted for ...",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")
    doc.add_paragraph(
        "Sem o admin consent, o sistema vai conseguir pedir token, mas nao vai conseguir enviar mensagens."
    )


def add_restriction_steps(doc: Document) -> None:
    doc.add_heading("7. Restringir o aplicativo para usar apenas a mailbox de envio", level=1)
    doc.add_paragraph(
        "Mesmo com Mail.Send, nao e boa pratica deixar a app enviar como qualquer usuario do tenant. "
        "Por isso, vamos restringir a app para a caixa recebimento@pibn.org.br."
    )
    doc.add_heading("7.1 Criar um grupo de seguranca com e-mail", level=2)
    for step in [
        "No Microsoft 365 admin center, crie um mail-enabled security group.",
        "Sugestao de nome: grp-powerchurch-mailer.",
        "Adicione recebimento@pibn.org.br como membro desse grupo.",
        "Anote o endereco do grupo, por exemplo: grp-powerchurch-mailer@seudominio.",
    ]:
        doc.add_paragraph(step, style="List Number")
    doc.add_heading("7.2 Aplicar a policy no Exchange Online PowerShell", level=2)
    doc.add_paragraph("Com PowerShell, execute na ordem:")
    code_blocks = [
        "Install-Module ExchangeOnlineManagement",
        "Connect-ExchangeOnline",
        'New-ApplicationAccessPolicy -AppId \"SEU-CLIENT-ID\" -PolicyScopeGroupId \"grp-powerchurch-mailer@seudominio\" -AccessRight RestrictAccess -Description \"Restringe o app Power Church a mailbox de recebimento\"',
        'Test-ApplicationAccessPolicy -AppId \"SEU-CLIENT-ID\" -Identity \"recebimento@pibn.org.br\"',
    ]
    for command in code_blocks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(command)
        r.font.name = "Courier New"
        r.font.size = Pt(10)
    doc.add_paragraph(
        "Se o teste indicar acesso permitido para essa mailbox, a restricao minima ficou pronta."
    )


def add_django_env(doc: Document) -> None:
    doc.add_heading("8. Configurar as variaveis de ambiente no Django", level=1)
    doc.add_paragraph(
        "No servidor do Power Church Django, vamos parar de depender de SMTP e passar a usar "
        "credenciais do Microsoft Graph."
    )
    env_lines = [
        "POWER_CHURCH_EMAIL_PROVIDER=microsoft_graph",
        "POWER_CHURCH_GRAPH_TENANT_ID=cole-aqui-o-tenant-id",
        "POWER_CHURCH_GRAPH_CLIENT_ID=cole-aqui-o-client-id",
        "POWER_CHURCH_GRAPH_CLIENT_SECRET=cole-aqui-o-client-secret",
        "POWER_CHURCH_GRAPH_SENDER_USER=recebimento@pibn.org.br",
    ]
    for line in env_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10.5)
    doc.add_paragraph(
        "Esses nomes podem ser ajustados no codigo, mas a recomendacao e padronizar assim para facilitar manutencao."
    )


def add_python_flow(doc: Document) -> None:
    doc.add_heading("9. O que o Django vai fazer internamente", level=1)
    doc.add_paragraph("O fluxo recomendado dentro do sistema e este:")
    flow = [
        "Gerar o PDF do recibo no backend.",
        "Ler o modelo padrao de assunto e corpo do e-mail.",
        "Pedir um token OAuth 2.0 ao Microsoft Entra usando client credentials.",
        "Montar a mensagem em formato Microsoft Graph.",
        "Anexar o PDF em base64.",
        "Enviar para /users/recebimento@pibn.org.br/sendMail.",
        "Registrar sucesso ou falha na fila ReceiptDispatch.",
    ]
    for step in flow:
        doc.add_paragraph(step, style="List Number")
    doc.add_heading("9.1 Exemplo de requisicao do token", level=2)
    token_example = [
        "POST https://login.microsoftonline.com/<TENANT_ID>/oauth2/v2.0/token",
        "grant_type=client_credentials",
        "client_id=<CLIENT_ID>",
        "client_secret=<CLIENT_SECRET>",
        "scope=https://graph.microsoft.com/.default",
    ]
    for line in token_example:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10)
    doc.add_heading("9.2 Exemplo de envio da mensagem", level=2)
    doc.add_paragraph("Endpoint:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("POST https://graph.microsoft.com/v1.0/users/recebimento@pibn.org.br/sendMail")
    r.font.name = "Courier New"
    r.font.size = Pt(10)


def add_json_example(doc: Document) -> None:
    doc.add_heading("10. Exemplo simplificado do JSON do envio", level=1)
    payload_lines = [
        "{",
        '  "message": {',
        '    "subject": "Recibo de contribuicoes - Maio 2026",',
        '    "body": {',
        '      "contentType": "Text",',
        '      "content": "Segue em anexo o seu recibo."',
        "    },",
        '    "toRecipients": [{"emailAddress": {"address": "destinatario@exemplo.com"}}],',
        '    "attachments": [',
        '      {',
        '        "@odata.type": "#microsoft.graph.fileAttachment",',
        '        "name": "recibo_maio_2026.pdf",',
        '        "contentType": "application/pdf",',
        '        "contentBytes": "<pdf-em-base64>"',
        "      }",
        "    ]",
        "  },",
        '  "saveToSentItems": true',
        "}",
    ]
    for line in payload_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9.5)


def add_test_plan(doc: Document) -> None:
    doc.add_heading("11. Como testar sem colocar tudo em producao de uma vez", level=1)
    tests = [
        "Primeiro teste: obter o token e verificar se vem access_token.",
        "Segundo teste: enviar um e-mail simples sem anexo para uma conta sua.",
        "Terceiro teste: enviar um e-mail com um PDF pequeno de teste.",
        "Quarto teste: usar a fila do Power Church para um unico recibo real.",
        "Quinto teste: validar se o e-mail chegou, se o remetente ficou correto e se o PDF abriu.",
        "Sexto teste: validar reenvio manual.",
        "Setimo teste: validar envio automatico ao lancar uma contribuicao de pessoa com e-mail.",
    ]
    for test in tests:
        doc.add_paragraph(test, style="List Number")
    doc.add_paragraph(
        "A ordem e importante. Nao pule direto para envio automatico antes de provar token, mensagem simples e anexo."
    )


def add_errors(doc: Document) -> None:
    doc.add_heading("12. Erros comuns e como interpretar", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for idx, text in enumerate(["Erro", "Significado / acao recomendada"]):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_width(cell, 2.1 if idx == 0 else 4.4)
        set_cell_shading(cell, "E8EEF5")
    entries = [
        ("401 Unauthorized", "Client ID, client secret ou tenant ID incorretos; confirme os tres."),
        ("403 Forbidden", "A app existe, mas nao recebeu consentimento ou nao tem acesso a mailbox."),
        ("ErrorAccessDenied", "Policy de acesso ou permissao do Graph bloqueando o envio."),
        ("InvalidAuthenticationToken", "Token expirado, escopo errado ou token mal montado."),
        ("MailboxNotEnabledForRESTAPI", "Conta de envio com algum bloqueio/estado inconsistente no Exchange Online."),
        ("Attachment size issue", "PDF grande demais; validar tamanho do anexo."),
    ]
    for left, right in entries:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        set_cell_width(row[0], 2.1)
        set_cell_width(row[1], 4.4)


def add_cutover(doc: Document) -> None:
    doc.add_heading("13. Plano de entrada em producao", level=1)
    plan = [
        "Fase 1: criar a app e validar token.",
        "Fase 2: validar envio simples com um e-mail seu.",
        "Fase 3: validar recibo real de uma pessoa em teste.",
        "Fase 4: ligar reenvio manual na tela de recibos.",
        "Fase 5: ligar envio automatico por evento apenas para homologacao.",
        "Fase 6: monitorar alguns dias.",
        "Fase 7: ativar para producao plena.",
    ]
    for item in plan:
        doc.add_paragraph(item, style="List Number")


def add_final_checklist(doc: Document) -> None:
    doc.add_heading("14. Checklist final do iniciante", level=1)
    checklist = [
        "Eu consegui criar a app no Entra ID.",
        "Eu anotei Tenant ID, Client ID e Client Secret.",
        "Eu concedi Mail.Send com admin consent.",
        "Eu restringi a app para a mailbox de envio.",
        "Eu configurei as variaveis no Django.",
        "Eu consegui gerar token.",
        "Eu consegui enviar um e-mail simples.",
        "Eu consegui enviar um recibo PDF.",
        "Eu validei reenvio manual.",
        "Eu validei envio automatico por evento.",
    ]
    for item in checklist:
        doc.add_paragraph(f"[   ] {item}", style="List Bullet")


def add_next_steps(doc: Document) -> None:
    doc.add_heading("15. Proximo passo no Power Church", level=1)
    doc.add_paragraph(
        "Depois que este roteiro estiver aprovado no tenant Microsoft, o proximo passo tecnico no projeto sera "
        "trocar o backend de envio do Power Church para Microsoft Graph, manter a fila ReceiptDispatch, "
        "preservar o modelo padrao editavel pelo operador e usar o PDF do recibo ja implementado no Django."
    )
    doc.add_paragraph(
        "Recomendacao final: guardar este documento junto com as credenciais, o nome da app registrada e "
        "o procedimento de renovacao do client secret."
    )


def build_doc() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = base_document()
    add_header_footer(doc)
    add_title(doc)
    add_intro_table(doc)
    add_prereqs(doc)
    add_checklist_table(doc)
    add_entra_steps(doc)
    add_secret_steps(doc)
    add_permissions_steps(doc)
    add_restriction_steps(doc)
    add_django_env(doc)
    add_python_flow(doc)
    add_json_example(doc)
    add_test_plan(doc)
    add_errors(doc)
    add_cutover(doc)
    add_final_checklist(doc)
    add_next_steps(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
