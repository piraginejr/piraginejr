from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/piraginejr/Documents/New project/Teste/Power Church")
OUTPUT = ROOT / "data" / "homologacao" / "ROTEIRO_OPERADOR_VALIDACAO_ETAPAS_1_E_2_V1.docx"


PRIMARY = RGBColor(27, 54, 93)
ACCENT = RGBColor(56, 102, 166)
LIGHT = RGBColor(236, 242, 250)
TEXT = RGBColor(34, 34, 34)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, *, bold=False, size=10.5, color=TEXT):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(5)
    pf.line_spacing = 1.15


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Roteiro do Operador", bold=True, size=22, color=PRIMARY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(
        p,
        "Validacao rapida da migracao local para Django + PostgreSQL — Etapas 1 e 2",
        size=12.5,
        color=ACCENT,
    )

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Cm(4.2), Cm(6.5), Cm(6.6)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    labels = [
        ("Ambiente", "Django local em PostgreSQL", "URL sugerida", "http://127.0.0.1:63621/"),
        ("Escopo", "Modulos modernos + cadastro/familias", "Operador", "Usuario administrativo habitual"),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(labels):
        for col_idx, value in enumerate([l1, v1, l2]):
            cell = table.rows[row_idx].cells[col_idx]
            set_cell_margins(cell)
            set_cell_shading(cell, "EDF2FA" if col_idx % 2 == 0 else "FFFFFF")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value, bold=col_idx % 2 == 0, size=10)
        last = table.rows[row_idx].cells[2]
        last.text = ""
        p = last.paragraphs[0]
        add_run(p, l2, bold=True, size=10)
        p = last.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        add_run(p, v2, size=10)
        set_cell_margins(last)
        set_cell_shading(last, "EDF2FA")

    doc.add_paragraph("")


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    add_run(p, text, bold=True, size=13, color=PRIMARY)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        add_run(p, "• ", bold=True, size=10.5, color=PRIMARY)
        add_run(p, item, size=10.5)


def add_steps_table(doc: Document, title: str, rows_data: list[tuple[str, str, str]]) -> None:
    add_section_heading(doc, title)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    headers = ["Passo", "Acao do operador", "Resultado esperado", "Status", "Observacao"]
    widths = [Cm(1.3), Cm(5.4), Cm(7.0), Cm(2.0), Cm(2.3)]
    for idx, width in enumerate(widths):
        table.rows[0].cells[idx].width = width
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "1B365D")
        set_cell_margins(cell, top=85, start=95, bottom=85, end=95)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_run(p, text, bold=True, size=9.5, color=RGBColor(255, 255, 255))
    for num, action, expected in rows_data:
        cells = table.add_row().cells
        values = [num, action, expected, "", ""]
        for idx, value in enumerate(values):
            cell = cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if idx in {3, 4}:
                set_cell_shading(cell, "F8FAFD")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_run(p, value, size=9.5)
    doc.add_paragraph("")


def add_acceptance_box(doc: Document, title: str, items: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "EDF2FA")
    set_cell_margins(cell, top=110, start=130, bottom=110, end=130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, bold=True, size=11, color=PRIMARY)
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(1)
        add_run(p, "• ", bold=True, size=10, color=PRIMARY)
        add_run(p, item, size=10)
    doc.add_paragraph("")


def build_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title_block(doc)

    add_section_heading(doc, "1. Como usar este roteiro")
    add_bullets(
        doc,
        [
            "Este documento consolida os testes rapidos que o operador deve executar para validar as mudancas concluidas ate agora na migracao.",
            "Antes de iniciar, a checagem tecnica das etapas deve estar verde.",
            "O operador deve marcar cada linha como OK, Falhou ou Observacao, sem tentar testar modulos fora do escopo abaixo.",
        ],
    )
    add_acceptance_box(
        doc,
        "Pre-condicao tecnica antes de liberar o teste humano",
        [
            'No projeto, rodar: python3 scripts/executar_homologacao_migracao.py --stage 1',
            'Depois, rodar: python3 scripts/executar_homologacao_migracao.py --stage 2',
            "Se houver falha tecnica, nao liberar o roteiro para o operador.",
        ],
    )

    add_section_heading(doc, "2. Escopo validado ate agora")
    add_bullets(
        doc,
        [
            "Etapa 1: fundacao do ambiente Django + PostgreSQL, navegacao principal, auditoria, recibos e monitor da fila.",
            "Etapa 2 (bloco concluido): lista de pessoas, busca por nome/ID/CPF/e-mail, ficha da pessoa, edicao cadastral, familias domiciliares, fila de auditoria familiar e merge auditavel.",
            "Importacoes bancarias e conciliacoes continuam fora do foco deste roteiro; elas serao validadas no bloco financeiro posterior.",
        ],
    )

    add_steps_table(
        doc,
        "3. Etapa 1 — Fundacao e modulos modernos",
        [
            ("1", "Abrir o sistema no ambiente Postgres", "Sistema abre sem erro e com layout completo."),
            ("2", "Fazer login com usuario valido", "Login aceito e dashboard carregado."),
            ("3", "Abrir Auditoria", "Tela abre com eventos e comandos visiveis."),
            ("4", "Abrir Recibos", "Central de recibos abre sem regressao visual."),
            ("5", "Abrir Monitor da fila", "Resumo, filtros e tags aparecem corretamente."),
            ("6", "Abrir um envio de e-mail na auditoria", "Detalhe mostra assunto, destinatario e conteudo."),
            ("7", "Navegar entre dashboard, auditoria e recibos", "Navegacao principal funciona sem links quebrados."),
            ("8", "Conferir se os comandos principais estao visiveis", "Botoes principais aparecem no lugar esperado."),
            ("9", "Validar se nao houve quebra de layout", "As telas continuam legiveis e operacionais."),
        ],
    )
    add_acceptance_box(
        doc,
        "Aceite da Etapa 1",
        [
            "Todos os passos da Etapa 1 marcados como OK.",
            "Sem regressao de navegacao.",
            "Auditoria, central de recibos e monitor da fila abrindo normalmente.",
        ],
    )

    add_steps_table(
        doc,
        "4. Etapa 2 — Cadastro, buscas, familias e merge",
        [
            ("1", "Buscar pessoa por nome", "Resultado correto aparece na lista."),
            ("2", "Buscar pessoa por ID ou codigo", "A ficha correta e localizada."),
            ("3", "Buscar por CPF", "Retorna a pessoa correta sem erro."),
            ("4", "Buscar por e-mail", "Lista localiza a pessoa correta."),
            ("5", "Abrir a ficha da pessoa", "Ficha carrega com dados cadastrais completos."),
            ("6", "Editar telefone, e-mail ou endereco de teste", "Gravacao funciona e dados reaparecem corretamente."),
            ("7", "Validar contatos e endereco apos salvar", "Campos ficam persistidos e coerentes."),
            ("8", "Abrir Familias domiciliares", "Lista carrega com filtros e contadores."),
            ("9", "Testar filtro de familias", "Resultado responde ao filtro sem erro visual."),
            ("10", "Abrir fila de auditoria das familias", "Sugestoes e grupos aparecem corretamente."),
            ("11", "Executar um merge controlado de caso conhecido", "Comparacao abre, justificativa grava e historico fica auditado."),
            ("12", "Confirmar historico/auditoria do merge", "Evento aparece com trilha consistente."),
        ],
    )
    add_acceptance_box(
        doc,
        "Aceite da Etapa 2",
        [
            "Buscas por nome, ID, CPF e e-mail marcadas como OK.",
            "Ficha, edicao, contatos e endereco marcados como OK.",
            "Familias e fila de auditoria abrindo corretamente.",
            "Merge controlado funcionando com historico auditavel.",
        ],
    )

    add_section_heading(doc, "5. Massa de prova sugerida")
    add_bullets(
        doc,
        [
            "Paschoal Piragine Junior.",
            "Um cadastro com e-mail valido.",
            "Um cadastro com CPF conhecido.",
            "Uma familia domiciliar conhecida para conferenciar agrupamento.",
            "Um caso de merge controlado previamente combinado para nao gerar risco operacional.",
        ],
    )

    add_section_heading(doc, "6. Observacoes finais do operador")
    add_bullets(
        doc,
        [
            "Se algum passo falhar, registrar exatamente em qual tela ocorreu e o que era esperado.",
            "Se houver divergencia visual, registrar se foi so layout ou se impactou a operacao.",
            "Nao avancar para o bloco financeiro da migracao usando este roteiro; ele cobre apenas o que foi concluido ate agora.",
        ],
    )

    section = doc.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(footer, "Power Church — roteiro de homologacao do operador", size=8.5, color=ACCENT)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_doc()
    print(path)
